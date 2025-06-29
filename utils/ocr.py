import easyocr
import tempfile
from langdetect import detect
from PIL import Image
from docx import Document
import os
import pdfplumber
from pdf2image import convert_from_path
from striprtf.striprtf import rtf_to_text
from odf.opendocument import load as load_odt
from odf.text import P, H
import csv
import re

def clean_text(text):
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    text = re.sub(r'[_~“”‘’¥£•—–…]', '', text)
    text = re.sub(r'-\s*\n', '', text)
    text = re.sub(r'\n+', '\n', text)
    text = re.sub(r'\s{2,}', ' ', text)
    return text.strip()

def clean_hyphens(text):
    return re.sub(r'(?<=\w)-\s+(?=\w)', '', text)

def extract_text(uploaded_file):
    suffix = os.path.splitext(uploaded_file.name)[-1].lower()
    text = ""

    if suffix in [".png", ".jpg", ".jpeg"]:
        image = Image.open(uploaded_file).convert('RGB')
        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp:
            image.save(tmp.name)
            reader = easyocr.Reader(['en'], gpu=False)
            result = reader.readtext(tmp.name, detail=0)
        text = "\n".join(result)

    elif suffix == ".pdf":
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(uploaded_file.read())
            tmp.flush()

            text_lines = []

            # Extract text from PDF using pdfplumber
            with pdfplumber.open(tmp.name) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        lines = page_text.splitlines()
                        text_lines.extend(lines)

            # OCR on images of PDF
            images = convert_from_path(tmp.name, dpi=300)
            reader = easyocr.Reader(['en'], gpu=False)
            for img in images:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as img_tmp:
                    img.save(img_tmp.name)
                    result = reader.readtext(img_tmp.name, detail=0)
                    text_lines.extend(result)

            # Deduplicate
            seen = set()
            unique_lines = []
            for line in text_lines:
                clean = line.strip()
                if clean and clean not in seen:
                    seen.add(clean)
                    unique_lines.append(clean)

            text = "\n".join(unique_lines)

    elif suffix == ".docx":
        doc = Document(uploaded_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + "\t".join(cell.text.strip() for cell in row.cells)

    elif suffix == ".txt":
        text = uploaded_file.read().decode("utf-8")

    elif suffix == ".rtf":
        text = rtf_to_text(uploaded_file.read().decode("utf-8"))

    elif suffix == ".csv":
        decoded = uploaded_file.read().decode("utf-8").splitlines()
        reader = csv.reader(decoded)
        lines = ["\t".join(row) for row in reader]
        text = "\n".join(lines)

    elif suffix == ".odt":
        with tempfile.NamedTemporaryFile(delete=False, suffix=".odt") as tmp:
            tmp.write(uploaded_file.read())
            tmp.flush()
            odt_doc = load_odt(tmp.name)

            paragraphs = odt_doc.getElementsByType(P) + odt_doc.getElementsByType(H)
            text_lines = []
            for p in paragraphs:
                try:
                    text_lines.append(str(p.firstChild.data))
                except AttributeError:
                    continue  # Skip empty or malformed elements
            text = "\n".join(text_lines)

    else:
        raise ValueError("Unsupported file type: " + suffix)

    text = clean_hyphens(text)
    text = clean_text(text)
    language = detect(text) if text.strip() else "unknown"
    return text.strip(), language
